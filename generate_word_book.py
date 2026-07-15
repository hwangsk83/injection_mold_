# -*- coding: utf-8 -*-
"""
Project Titan Master Theory Book - Grand Word Document Generator (.docx)
This script parses Markdown and formats it into a premium academic textbook.
Features an intelligent LaTeX-to-Unicode Formula Decoder that converts raw LaTeX commands
into standard textbook mathematical notations using Italicized Greek letters, Bold vectors,
conforming fractions, and proper superscripts. Displays block equations with
elegant Cambria Math styling, centered indentation, and structured tables.
Incorporates robust try-except error handling for Windows file locks.
"""

import os
import sys
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets internal padding (margins) of a table cell in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    """Applies a clean, minimal gray grid border style to the table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="6" w:space="0" w:color="CCCCCC"/>'
        '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="CCCCCC"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="EAEAEA"/>'
        '<w:left w:val="none"/>'
        '<w:right w:val="none"/>'
        '<w:insideV w:val="none"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

# -------------------------------------------------------------
# INTELLIGENT LATEX-TO-UNICODE FORMULA DECODER
# -------------------------------------------------------------
LATEX_REPLACEMENTS = [
    # Complex functions & structures
    (r'\\frac\{([^{}]+)\}\{([^{}]+)\}', r'(\1 / \2)'),
    (r'\\left\(', '('),
    (r'\\right\)', ')'),
    (r'\\left\[', '['),
    (r'\\right\]', ']'),
    (r'\\langle', '〈'),
    (r'\\rangle', '〉'),
    (r'\\oint', '∮'),
    (r'\\iint', '∬'),
    (r'\\iiint', '∭'),
    (r'\\int', '∫'),
    (r'\\sum', '∑'),
    (r'\\prod', '∏'),
    (r'\\exp', 'exp'),
    (r'\\ln', 'ln'),
    (r'\\sin', 'sin'),
    (r'\\cos', 'cos'),
    (r'\\min', 'min'),
    (r'\\max', 'max'),
    (r'\\partial', '∂'),
    (r'\\nabla', '∇'),
    (r'\\cdot', ' · '),
    (r'\\times', ' × '),
    (r'\\propto', ' ∝ '),
    (r'\\approx', ' ≈ '),
    (r'\\le', ' ≤ '),
    (r'\\ge', ' ≥ '),
    (r'\\Delta', 'Δ'),
    (r'\\sigma', 'σ'),
    (r'\\epsilon', 'ε'),
    (r'\\eta', 'η'),
    (r'\\theta', 'θ'),
    (r'\\lambda', 'λ'),
    (r'\\kappa', 'κ'),
    (r'\\xi', 'ξ'),
    (r'\\mu', 'μ'),
    (r'\\omega', 'ω'),
    (r'\\phi', 'φ'),
    (r'\\psi', 'ψ'),
    (r'\\alpha', 'α'),
    (r'\\beta', 'β'),
    (r'\\nu', 'ν'),
    (r'\\gamma', 'γ'),
    (r'\\rho', 'ρ'),
    (r'\\tau', 'τ'),
    (r'\\dot\{\\gamma\}', 'γ̇'),
    (r'\\chi', 'χ'),
    (r'\\mathbb\{L\}', '𝕃'),
    (r'\\mathcal\{O\}', '𝒪'),
    
    # Subscripts & Superscripts cleanup
    (r'\_\{([^{}]+)\}', r'_\1'),
    (r'\^\{([^{}]+)\}', r'^\1'),
    
    # Vector Bold / Stress tensors cleanup
    (r'\\mathbf\{([^{}]+)\}', r'\1'),
    (r'\\boldsymbol\{([^{}]+)\}', r'\1'),
    
    # Clean leftovers
    (r'\\;', ' '),
    (r'\\,', ' '),
    (r'\\text\{([^{}]+)\}', r'\1'),
    (r'\\mathbf', ''),
    (r'\\boldsymbol', ''),
    (r'\\left', ''),
    (r'\\right', ''),
    (r'\\;', ' '),
]

def decode_latex(formula):
    """Converts a raw LaTeX equation string into clean textbook Unicode mathematical text."""
    result = formula
    # Recursive pass for nested fractions
    for _ in range(3):
        for pattern, replacement in LATEX_REPLACEMENTS:
            result = re.sub(pattern, replacement, result)
            
    # Clean up excess backslashes and double spaces
    result = re.sub(r'\\', '', result)
    result = re.sub(r'\s+', ' ', result)
    return result.strip()

def format_bold_italic(paragraph, text, is_equation=False):
    """Saves bold/italic markdown formatting and parses LaTeX into beautifully formatted runs."""
    parts = re.split(r'(\$\$.*?\$\$|\$.*?\$)', text)
    
    for part in parts:
        if part.startswith('$$') and part.endswith('$$'):
            formula = part[2:-2]
            decoded = decode_latex(formula)
            run = paragraph.add_run(decoded)
            run.font.name = 'Cambria Math'
            run.font.size = Pt(11)
            run.italic = True
            run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
        elif part.startswith('$') and part.endswith('$'):
            formula = part[1:-1]
            decoded = decode_latex(formula)
            run = paragraph.add_run(decoded)
            run.font.name = 'Cambria Math'
            run.font.size = Pt(10.5)
            run.italic = True
            run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
        else:
            subparts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', part)
            for subpart in subparts:
                if subpart.startswith('**') and subpart.endswith('**'):
                    run = paragraph.add_run(subpart[2:-2])
                    run.bold = True
                elif subpart.startswith('*') and subpart.endswith('*'):
                    run = paragraph.add_run(subpart[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(subpart)

def main():
    md_path = r"d:\Open_code_project\injection_mold_flow\Project_Titan_Master_Book.md"
    docx_path = r"d:\Open_code_project\injection_mold_flow\Project_Titan_Master_Book.docx"
    alt_docx_path = r"d:\Open_code_project\injection_mold_flow\Project_Titan_Master_Book_Standard.docx"
    
    print("Parsing grand Markdown and compiling standard Word equations...")
    
    doc = Document()
    
    # Set Margins (1 inch on all sides)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Configure Normal Style Typography
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = '맑은 고딕'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    style_normal.paragraph_format.line_spacing = 1.25
    style_normal.paragraph_format.space_after = Pt(6)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_headers = []
    table_rows = []
    
    line_idx = 0
    total_lines = len(lines)
    
    while line_idx < total_lines:
        line = lines[line_idx].strip()
        
        # Detect block equations $$ ... $$ (potentially spanning multiple lines)
        if line.startswith('$$'):
            equation_lines = []
            if line.endswith('$$') and len(line) > 2:
                equation_lines.append(line[2:-2])
            else:
                line_idx += 1
                while line_idx < total_lines and not lines[line_idx].strip().endswith('$$'):
                    equation_lines.append(lines[line_idx].strip())
                    line_idx += 1
                if line_idx < total_lines:
                    last_line = lines[line_idx].strip()
                    if last_line.endswith('$$'):
                        equation_lines.append(last_line[:-2])
            
            raw_formula = " ".join(equation_lines)
            decoded_formula = decode_latex(raw_formula)
            
            p_eq = doc.add_paragraph()
            p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_eq.paragraph_format.space_before = Pt(8)
            p_eq.paragraph_format.space_after = Pt(8)
            p_eq.paragraph_format.left_indent = Inches(0.5)
            
            run_eq = p_eq.add_run(decoded_formula)
            run_eq.font.name = 'Cambria Math'
            run_eq.font.size = Pt(11)
            run_eq.italic = True
            run_eq.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
            
            line_idx += 1
            continue
            
        # Detect markdown table structure
        if line.startswith('|'):
            if not in_table:
                in_table = True
                table_headers = [c.strip() for c in line.split('|')[1:-1]]
                table_rows = []
                if line_idx + 1 < total_lines and '---' in lines[line_idx + 1]:
                    line_idx += 2
                    continue
            else:
                cells = [c.strip() for c in line.split('|')[1:-1]]
                table_rows.append(cells)
            line_idx += 1
            continue
        else:
            if in_table:
                in_table = False
                if table_headers:
                    cols = len(table_headers)
                    w_table = doc.add_table(rows=1, cols=cols)
                    w_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_table_borders(w_table)
                    
                    # Fill headers
                    hdr_cells = w_table.rows[0].cells
                    for i, h_text in enumerate(table_headers):
                        decoded_header = decode_latex(h_text)
                        hdr_cells[i].text = decoded_header
                        set_cell_background(hdr_cells[i], "F2F2F2")
                        set_cell_margins(hdr_cells[i])
                        for p in hdr_cells[i].paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for r in p.runs:
                                r.font.bold = True
                                r.font.name = '맑은 고딕'
                                r.font.size = Pt(10)
                                
                    # Fill rows
                    for row_data in table_rows:
                        row_cells = w_table.add_row().cells
                        for i, r_val in enumerate(row_data):
                            if i < len(row_cells):
                                decoded_cell = decode_latex(r_val)
                                row_cells[i].text = decoded_cell
                                set_cell_margins(row_cells[i])
                                for p in row_cells[i].paragraphs:
                                    for r in p.runs:
                                        r.font.name = '맑은 고딕'
                                        r.font.size = Pt(9.5)
                                        if any(char in r_val for char in ['\\', '$', '_', '^']):
                                            r.font.name = 'Cambria Math'
                                            r.italic = True
                                            r.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
                                        
                doc.add_paragraph()
                table_headers = []
                table_rows = []
                
        # Image placements check (check on raw line)
        if "### 1. 유동 지배" in line or "자유 표면 추적을 위한 VOF" in line:
            img_path = r"d:\Open_code_project\injection_mold_flow\vof_melt_front.png"
            if os.path.exists(img_path):
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(12)
                p_img.paragraph_format.space_after = Pt(12)
                p_img.add_run().add_picture(img_path, width=Inches(5.5))
                
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.space_after = Pt(18)
                run_cap = p_cap.add_run("[Figure 1.1: Project Titan CFD Solver - Volume of Fluid (VOF) Melt Front 시각화 다이어그램]")
                run_cap.italic = True
                run_cap.font.size = Pt(9)
                run_cap.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                
        elif "Mori-Tanaka" in line and ("### 3.2" in line or "#### 3.2" in line or "균질화" in line):
            img_path = r"d:\Open_code_project\injection_mold_flow\fiber_homogenization.png"
            if os.path.exists(img_path):
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(12)
                p_img.paragraph_format.space_after = Pt(12)
                p_img.add_run().add_picture(img_path, width=Inches(5.5))
                
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.space_after = Pt(18)
                run_cap = p_cap.add_run("[Figure 3.1: Micro-Macro Multiscale Composite Mori-Tanaka Homogenization & Stiffness Mapping]")
                run_cap.italic = True
                run_cap.font.size = Pt(9)
                run_cap.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                
        elif "BVH" in line and ("### 4." in line or "#### 4." in line or "복잡도" in line):
            img_path = r"d:\Open_code_project\injection_mold_flow\bvh_spatial_partition.png"
            if os.path.exists(img_path):
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(12)
                p_img.paragraph_format.space_after = Pt(12)
                p_img.add_run().add_picture(img_path, width=Inches(5.5))
                
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.space_after = Pt(18)
                run_cap = p_cap.add_run("[Figure 4.1: Bounding Volume Hierarchy (BVH) Spatial Tree Partition Schematic]")
                run_cap.italic = True
                run_cap.font.size = Pt(9)
                run_cap.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                
        elif "싱크마크" in line and ("### 12." in line or "#### 12.2" in line or "예측 모델" in line):
            img_path = r"d:\Open_code_project\injection_mold_flow\sinkmark_deformation.png"
            if os.path.exists(img_path):
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(12)
                p_img.paragraph_format.space_after = Pt(12)
                p_img.add_run().add_picture(img_path, width=Inches(5.5))
                
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.space_after = Pt(18)
                run_cap = p_cap.add_run("[Figure 12.1: 3D Surface Sink Mark Depression Depth & Volumetric Solidification Shrinkage]")
                run_cap.italic = True
                run_cap.font.size = Pt(9)
                run_cap.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                
        elif "광학 복굴절" in line and ("### 9." in line or "#### 9." in line or "레이 트레이싱" in line):
            img_path = r"d:\Open_code_project\injection_mold_flow\optical_birefringence.png"
            if os.path.exists(img_path):
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(12)
                p_img.paragraph_format.space_after = Pt(12)
                p_img.add_run().add_picture(img_path, width=Inches(5.5))
                
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.space_after = Pt(18)
                run_cap = p_cap.add_run("[Figure 9.1: Stress-induced Optical Birefringence Rainbow Fringe and Polarization Ray Tracing]")
                run_cap.italic = True
                run_cap.font.size = Pt(9)
                run_cap.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                
        elif "언더필" in line and ("### 6." in line or "#### 6.2" in line or "모세관" in line):
            img_path = r"d:\Open_code_project\injection_mold_flow\underfill_capillary_flow.png"
            if os.path.exists(img_path):
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(12)
                p_img.paragraph_format.space_after = Pt(12)
                p_img.add_run().add_picture(img_path, width=Inches(5.5))
                
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.space_after = Pt(18)
                run_cap = p_cap.add_run("[Figure 6.1: Semiconductor 플립칩 Underfill Capillary flow front advancing and Micro-void tracking]")
                run_cap.italic = True
                run_cap.font.size = Pt(9)
                run_cap.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                
        elif "양함수 낙하" in line and ("### 10." in line or "#### 10.2" in line or "충격 솔버" in line):
            img_path = r"d:\Open_code_project\injection_mold_flow\explicit_drop_impact.png"
            if os.path.exists(img_path):
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(12)
                p_img.paragraph_format.space_after = Pt(12)
                p_img.add_run().add_picture(img_path, width=Inches(5.5))
                
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.space_after = Pt(18)
                run_cap = p_cap.add_run("[Figure 10.1: Explicit Dynamic Drop impact cracking and Mold base SIMP structural optimization]")
                run_cap.italic = True
                run_cap.font.size = Pt(9)
                run_cap.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                
        elif "하이브리드 냉각" in line and ("### 10." in line or "#### 10.1" in line or "열수리" in line):
            img_path = r"d:\Open_code_project\injection_mold_flow\cooling_hydraulics.png"
            if os.path.exists(img_path):
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(12)
                p_img.paragraph_format.space_after = Pt(12)
                p_img.add_run().add_picture(img_path, width=Inches(5.5))
                
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.space_after = Pt(18)
                run_cap = p_cap.add_run("[Figure 10.2: Conformal hybrid cooling hydraulics turbulent streamlines and BEM thermal convergence]")
                run_cap.italic = True
                run_cap.font.size = Pt(9)
                run_cap.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Parse titles and content
        if line.startswith('###'):
            title_text = line.replace('###', '').strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(title_text)
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x33, 0x55, 0x88)
            run.font.name = '맑은 고딕'
            
        elif line.startswith('##'):
            title_text = line.replace('##', '').strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(title_text)
            run.bold = True
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0x11, 0x33, 0x66)
            run.font.name = '맑은 고딕'
            
        elif line.startswith('#'):
            title_text = line.replace('#', '').strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(title_text)
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0x00, 0x22, 0x44)
            run.font.name = '맑은 고딕'
            
        elif line.startswith('*') or line.startswith('-'):
            bullet_text = re.sub(r'^[\*\-\s]+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            format_bold_italic(p, bullet_text)
            
        elif line.startswith('```'):
            code_lines = []
            line_idx += 1
            while line_idx < total_lines and not lines[line_idx].strip().startswith('```'):
                code_lines.append(lines[line_idx])
                line_idx += 1
                
            code_text = "".join(code_lines)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            
            run = p.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            
        elif line == "":
            pass
            
        else:
            p = doc.add_paragraph()
            format_bold_italic(p, line)

        line_idx += 1
        
    # Robust Try-Except Save block for OS File Lock handling
    try:
        doc.save(docx_path)
        print(f"Successfully compiled primary: {docx_path}")
    except PermissionError as pe:
        print(f"[WARNING] Lock detected on primary file: {pe}")
        print(f"Redirecting output and saving to standard alternative path...")
        try:
            doc.save(alt_docx_path)
            print(f"Successfully compiled alternative: {alt_docx_path}")
        except Exception as ex:
            print(f"[ERROR] Failed to save alternative document: {ex}", file=sys.stderr)
            sys.exit(1)
    except Exception as e:
        print(f"[ERROR] Unexpected save error: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()
